# -*- coding: utf-8 -*-
"""
维权通数据转换脚本
功能：Word文档 → Excel(维护层) → JSON(运行层)
支持：首次提取、增量更新(diff)、数据校验、一键转换

用法：
  # 首次提取：Word → Excel + JSON
  python convert.py --input "./docs" --output "./data" --mode all

  # 仅提取Word → Excel
  python convert.py --input "./docs" --output "./data/channels.xlsx" --mode extract

  # 仅构建Excel → JSON
  python convert.py --input "./data/channels.xlsx" --output "./miniprogram/data" --mode build

  # 仅校验
  python convert.py --input "./data/channels.xlsx" --mode validate

  # 增量更新（与现有Excel对比，输出diff报告）
  python convert.py --input "./docs" --output "./data/new.xlsx" --existing "./data/current.xlsx" --diff

依赖：pip install python-docx openpyxl jsonschema tqdm
"""

import os
import re
import json
import hashlib
import argparse
import logging
from datetime import datetime
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print("缺少依赖 python-docx，请运行: pip install python-docx")
    exit(1)

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
except ImportError:
    print("缺少依赖 openpyxl，请运行: pip install openpyxl")
    exit(1)

try:
    from tqdm import tqdm
except ImportError:
    # tqdm可选，没有则用简单进度显示
    def tqdm(iterable, **kwargs):
        return iterable

# ============================================================
# 配置常量
# ============================================================

# 渠道表格的键字段映射（Word中的键名 → 标准字段名）
CHANNEL_FIELD_MAP = {
    '渠道名称': 'name',
    '平台名称': 'name',
    '官方网址': 'website',
    '投诉电话': 'phone',
    '联系电话': 'phone',
    '上级监管部门': 'regulator',
    '监管部门': 'regulator',
    '适用范围': 'scope',
    '受理范围': 'scope',
    '信息来源': 'source',
    '法律依据': 'legal_basis',
    '实用提示': 'tips',
    '威慑力提示': 'tips',
    '前置条件': 'precondition',
    '注意事项': 'precondition',
    '电话备注': 'phone_note',
    '服务时间': 'phone_note',
}

# 一级分类名称（用于识别Heading 1）
CATEGORY_L1_NAMES = [
    '基础民生与公共交通',
    '金融与商业消费',
    '社会服务与政务司法',
    '高层级诉求平台',
    '群众诉求官方平台',
    '投诉求助类',
    '问政建议类',
    '四川省',
    '四川/成都地方渠道',
    '实战工具与模板',
    '投诉实用指南',
    '投诉效力真相',
]

# 城市编码映射
CITY_CODE_MAP = {
    '成都': 'chengdu',
    '四川': 'sichuan',
    '全国': 'national',
    '国家': 'national',
}

# 枚举值
VALID_STATUS = ['active', 'merged', 'discontinued']
VALID_CITY_CODE = ['national', 'sichuan', 'chengdu']
VALID_CHANNEL_TYPE = ['official', 'hotline', 'platform', 'enterprise']

# 渠道表Excel列定义（字段名, 中文说明, 是否必填, 列宽）
CHANNEL_COLUMNS = [
    ('id', '渠道ID', True, 12),
    ('name', '渠道名称', True, 35),
    ('phone', '投诉电话', False, 18),
    ('phone_note', '电话备注', False, 25),
    ('website', '官方网址', False, 45),
    ('regulator', '上级监管部门', False, 25),
    ('scope', '适用范围', True, 60),
    ('precondition', '前置条件', False, 40),
    ('legal_basis', '法律依据', False, 40),
    ('tips', '实用提示', False, 40),
    ('source', '信息来源', False, 35),
    ('category_l1', '一级分类', True, 22),
    ('category_l2', '二级分类', True, 22),
    ('tags', '搜索标签(逗号分隔)', True, 40),
    ('city_code', '城市编码', True, 12),
    ('channel_type', '渠道类型', True, 14),
    ('status', '状态', True, 14),
    ('merged_to', '替代渠道ID', False, 12),
    ('hot_level', '热度权重(1-5)', True, 12),
    ('related_script_id', '关联话术ID', False, 12),
    ('law_ids', '关联法律法规ID(逗号分隔)', False, 30),
    ('effect_rating', '效力评级(V2.0)', False, 14),
    ('response_speed', '响应速度(V2.0)', False, 14),
    ('penalty_power', '处罚力度(V2.0)', False, 14),
    ('success_rate', '实测成功率(V2.5)', False, 14),
    ('user_feedback_count', '反馈数(V2.5)', False, 12),
    ('ext', '扩展字段(JSON)', False, 30),
]

# 话术表Excel列定义
SCRIPT_COLUMNS = [
    ('id', '话术ID', True, 12),
    ('scene_name', '场景名称', True, 30),
    ('applicable', '适用情况', True, 50),
    ('related_channel_id', '关联渠道ID', False, 12),
    ('script_type', '话术类型', True, 16),
    ('phone_script', '电话版话术', False, 80),
    ('written_complainant', '书面版-投诉人', False, 30),
    ('written_respondent', '书面版-被投诉人', False, 30),
    ('written_request', '书面版-投诉请求', False, 50),
    ('written_facts', '书面版-事实与理由', False, 80),
    ('written_evidence', '书面版-证据清单', False, 40),
    ('legal_basis', '法律依据', False, 40),
    ('law_ids', '关联法律法规ID(逗号分隔)', False, 30),
    ('evidence_list', '证据清单(逗号分隔)', False, 40),
    ('keywords', '搜索关键词(逗号分隔)', True, 40),
    ('is_hot', '是否首页展示', True, 10),
    ('hot_level', '热度权重(1-5)', True, 12),
    ('ext', '扩展字段(JSON)', False, 30),
]

# 配置表
DEFAULT_CONFIG = {
    "data_version": "2026.08",
    "data_verified_at": "2026-08-30",
    "search_weights": {
        "name_exact": 100, "phone_exact": 90, "tag_exact": 80,
        "scope_fuzzy": 50, "script_match": 60, "platform_match": 55
    },
    "hot_search_words": [
        "快递丢失", "运营商扣费", "商家不退款", "物业乱收费", "欠薪",
        "电信诈骗", "噪音扰民", "医院乱收费", "出租车拒载", "个人信息泄露"
    ],
    "synonyms": {
        "丢了": "丢失", "弄丢": "丢失", "坑人": "误导", "骗人": "欺诈",
        "不管": "不作为", "乱收费": "违规收费", "退款": "退款",
        "假货": "假冒伪劣", "噪音": "噪声扰民", "欠薪": "拖欠工资",
        "诈骗": "电信网络诈骗"
    },
    "feature_flags": {
        "ai_assistant": False, "complaint_ticket": False, "evidence_manager": False,
        "community": False, "user_login": False, "effect_rating": False,
        "script_fill": False, "multi_city": False, "payment": False
    },
    "emergency_phones": [
        {"name": "110报警", "phone": "110", "icon": "🚨"},
        {"name": "119火警", "phone": "119", "icon": "🔥"},
        {"name": "120急救", "phone": "120", "icon": "🚑"},
        {"name": "96110反诈", "phone": "96110", "icon": "🛡️"}
    ],
    "limits": {
        "view_history_max": 20, "search_history_max": 20,
        "search_suggest_max": 8, "search_page_size": 20
    }
}

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('weiquantong')


# ============================================================
# Word解析器
# ============================================================

class WordParser:
    """解析Word文档，提取渠道、话术、平台数据"""

    def __init__(self):
        self.channels = []
        self.scripts = []
        self.platforms = []
        self.warnings = []
        self._current_l1 = ''
        self._current_l2 = ''
        self._current_l3 = ''
        self._channel_id_counter = 0
        self._script_id_counter = 0
        self._platform_id_counter = 0

    def parse_directory(self, dir_path):
        """解析目录下所有docx文件"""
        docx_files = []
        for root, dirs, files in os.walk(dir_path):
            for f in files:
                if f.endswith('.docx') and not f.startswith('~$'):
                    docx_files.append(os.path.join(root, f))

        if not docx_files:
            logger.error(f"目录 {dir_path} 下未找到docx文件")
            return

        logger.info(f"找到 {len(docx_files)} 个Word文档")
        for filepath in sorted(docx_files):
            logger.info(f"解析: {os.path.basename(filepath)}")
            self.parse_file(filepath)

        # 去重
        self._deduplicate()
        logger.info(f"解析完成: 渠道{len(self.channels)}条, 话术{len(self.scripts)}条, 平台{len(self.platforms)}条")

    def parse_file(self, filepath):
        """解析单个Word文件"""
        try:
            doc = Document(filepath)
        except Exception as e:
            self.warnings.append(f"文件解析失败: {filepath}, 错误: {e}")
            return

        # 按文档顺序遍历段落和表格
        body = doc.element.body
        para_idx = 0
        table_idx = 0
        paragraphs = doc.paragraphs
        tables = doc.tables

        for child in body.iterchildren():
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                if para_idx < len(paragraphs):
                    self._process_paragraph(paragraphs[para_idx])
                    para_idx += 1
            elif tag == 'tbl':
                if table_idx < len(tables):
                    self._process_table(tables[table_idx])
                    table_idx += 1

    def _process_paragraph(self, para):
        """处理段落，更新分类上下文"""
        text = para.text.strip()
        if not text:
            return

        style = para.style.name if para.style else ''

        # 识别标题层级
        if style.startswith('Heading 1') or (text.startswith('# ') and not text.startswith('## ')):
            clean_text = text.lstrip('# ').strip()
            # 检查是否是已知的一级分类
            for cat_name in CATEGORY_L1_NAMES:
                if cat_name in clean_text:
                    self._current_l1 = cat_name
                    self._current_l2 = ''
                    self._current_l3 = ''
                    return
            # 其他Heading 1也记录
            self._current_l1 = clean_text[:30]
            self._current_l2 = ''

        elif style.startswith('Heading 2') or text.startswith('## '):
            clean_text = text.lstrip('# ').strip()
            # 去除序号前缀（如"一、""1."）
            clean_text = re.sub(r'^[一二三四五六七八九十]+[、.．]\s*', '', clean_text)
            clean_text = re.sub(r'^\d+[、.．)\s]\s*', '', clean_text)
            self._current_l2 = clean_text[:30]
            self._current_l3 = ''

        elif style.startswith('Heading 3') or text.startswith('### '):
            clean_text = text.lstrip('# ').strip()
            clean_text = re.sub(r'^\d+[、.．)\s]\s*', '', clean_text)
            self._current_l3 = clean_text[:30]

    def _process_table(self, table):
        """处理表格，判断是渠道表还是话术表"""
        if len(table.rows) == 0:
            return

        first_cell = table.rows[0].cells[0].text.strip() if table.rows[0].cells else ''

        # 判断是否为渠道表（2列，第一列含"渠道名称"或"平台名称"）
        if len(table.columns) == 2 and ('渠道名称' in first_cell or '平台名称' in first_cell):
            self._parse_channel_table(table)
            return

        # 判断是否为话术表（单列表格，内容含"你好，我要"或"投诉人"）
        if len(table.columns) == 1 and len(table.rows) >= 1:
            cell_text = table.rows[0].cells[0].text.strip()
            if cell_text.startswith('"') and ('我要' in cell_text or '投诉' in cell_text):
                self._parse_script_table(table, 'phone')
                return
            if '投诉人' in cell_text or '被投诉人' in cell_text or '投诉请求' in cell_text:
                self._parse_script_table(table, 'written')
                return

        # 其他表格记录警告（供人工审核）
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            preview = first_cell[:30]
            self.warnings.append(f"未识别表格(首列: {preview}), 行{len(table.rows)}列{len(table.columns)}")

    def _parse_channel_table(self, table):
        """解析渠道表格"""
        channel = {}
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            if not key:
                continue
            # 映射字段名
            field_name = CHANNEL_FIELD_MAP.get(key)
            if field_name:
                # 同一字段多次出现时合并
                if field_name in channel and channel[field_name]:
                    channel[field_name] += ' / ' + value
                else:
                    channel[field_name] = value
            else:
                # 未映射的字段放入ext
                if 'ext' not in channel:
                    channel['ext'] = {}
                channel['ext'][key] = value

        if not channel.get('name') or not channel.get('scope'):
            self.warnings.append(f"渠道表缺少必填字段: name={channel.get('name', '无')}, scope={'有' if channel.get('scope') else '无'}")
            return

        # 自动分配ID
        self._channel_id_counter += 1
        channel['id'] = f'ch_{self._channel_id_counter:03d}'

        # 补充分类
        channel['category_l1'] = self._current_l1 or '未分类'
        channel['category_l2'] = self._current_l2 or self._current_l3 or '未分类'

        # 自动生成标签
        channel['tags'] = self._generate_tags(channel)

        # 推断城市编码
        channel['city_code'] = self._infer_city_code(channel)

        # 推断渠道类型
        channel['channel_type'] = self._infer_channel_type(channel)

        # 状态默认active
        channel['status'] = 'active'

        # 热度默认3
        channel['hot_level'] = 3

        # ext转JSON字符串
        if 'ext' in channel and channel['ext']:
            channel['ext'] = json.dumps(channel['ext'], ensure_ascii=False)
        else:
            channel['ext'] = ''

        self.channels.append(channel)

    def _parse_script_table(self, table, script_part):
        """解析话术表格（phone=电话版, written=书面版）"""
        text = '\n'.join(row.cells[0].text.strip() for row in table.rows if row.cells)

        # 查找最近的话术场景（基于当前分类上下文）
        scene_name = self._current_l3 or self._current_l2 or '未命名场景'

        # 查找是否已有该场景的话术
        existing = None
        for s in self.scripts:
            if s['scene_name'] == scene_name:
                existing = s
                break

        if script_part == 'phone':
            if existing:
                existing['phone_script'] = text
            else:
                self._script_id_counter += 1
                script = {
                    'id': f'sc_{self._script_id_counter:03d}',
                    'scene_name': scene_name,
                    'applicable': '',
                    'related_channel_id': '',
                    'script_type': 'template',
                    'phone_script': text,
                    'written_complainant': '',
                    'written_respondent': '',
                    'written_request': '',
                    'written_facts': '',
                    'written_evidence': '',
                    'legal_basis': '',
                    'evidence_list': [],
                    'keywords': self._generate_script_keywords(scene_name),
                    'is_hot': True,
                    'hot_level': 4,
                    'ext': ''
                }
                self.scripts.append(script)
        elif script_part == 'written':
            if existing:
                # 解析书面版各部分
                self._parse_written_script(existing, text)
            else:
                self._script_id_counter += 1
                script = {
                    'id': f'sc_{self._script_id_counter:03d}',
                    'scene_name': scene_name,
                    'applicable': '',
                    'related_channel_id': '',
                    'script_type': 'template',
                    'phone_script': '',
                    'written_complainant': '',
                    'written_respondent': '',
                    'written_request': '',
                    'written_facts': '',
                    'written_evidence': '',
                    'legal_basis': '',
                    'evidence_list': [],
                    'keywords': self._generate_script_keywords(scene_name),
                    'is_hot': True,
                    'hot_level': 4,
                    'ext': ''
                }
                self._parse_written_script(script, text)
                self.scripts.append(script)

    def _parse_written_script(self, script, text):
        """解析书面版话术，拆分为各部分"""
        # 简单按关键词拆分
        if '投诉人' in text and not script['written_complainant']:
            match = re.search(r'投诉人[：:]\s*(.+?)(?=\n|被投诉人|投诉请求|$)', text, re.DOTALL)
            if match:
                script['written_complainant'] = match.group(1).strip()

        if '被投诉人' in text and not script['written_respondent']:
            match = re.search(r'被投诉人[：:]\s*(.+?)(?=\n|投诉请求|事实与理由|$)', text, re.DOTALL)
            if match:
                script['written_respondent'] = match.group(1).strip()

        if '投诉请求' in text and not script['written_request']:
            match = re.search(r'投诉请求[：:]\s*(.+?)(?=\n|事实与理由|证据清单|$)', text, re.DOTALL)
            if match:
                script['written_request'] = match.group(1).strip()

        if '事实与理由' in text and not script['written_facts']:
            match = re.search(r'事实与理由[：:]\s*(.+?)(?=\n|证据清单|综上|$)', text, re.DOTALL)
            if match:
                script['written_facts'] = match.group(1).strip()

        if '证据清单' in text and not script['written_evidence']:
            match = re.search(r'证据清单[：:]\s*(.+?)(?=\n|综上|此致|$)', text, re.DOTALL)
            if match:
                script['written_evidence'] = match.group(1).strip()

        # 如果拆分失败，全部放入written_facts
        if not script['written_facts'] and text:
            script['written_facts'] = text[:2000]

    def _generate_tags(self, channel):
        """自动生成搜索标签"""
        tags = set()

        # 从名称提取
        name = channel.get('name', '')
        for word in re.findall(r'[\u4e00-\u9fa5]{2,}', name):
            if len(word) <= 6:
                tags.add(word)

        # 从分类提取
        if channel.get('category_l2'):
            tags.add(channel['category_l2'])

        # 从适用范围提取高频词
        scope = channel.get('scope', '')
        high_freq_words = ['投诉', '举报', '申诉', '纠纷', '违法', '违规', '收费', '退款',
                           '丢失', '破损', '延误', '诈骗', '欺诈', '误导', '不作为', '乱收费',
                           '欠薪', '社保', '公积金', '物业', '快递', '电信', '银行', '保险',
                           '医疗', '教育', '食品', '药品', '环保', '劳动', '税务', '公安']
        for word in high_freq_words:
            if word in scope:
                tags.add(word)

        # 电话号码作为标签
        if channel.get('phone'):
            tags.add(channel['phone'])

        # 至少3个标签
        result = list(tags)[:8]
        if len(result) < 3:
            result.extend(['维权', '投诉', '官方渠道'])

        return result

    def _generate_script_keywords(self, scene_name):
        """生成话术搜索关键词"""
        keywords = [scene_name]
        if '快递' in scene_name:
            keywords.extend(['快递丢了', '快递破损', '快递慢', '12305'])
        if '运营商' in scene_name or '电信' in scene_name:
            keywords.extend(['话费乱扣', '运营商坑人', '12300'])
        if '商家' in scene_name or '消费' in scene_name:
            keywords.extend(['不退款', '假货', '虚假宣传', '12315'])
        if '银行' in scene_name or '保险' in scene_name or '金融' in scene_name:
            keywords.extend(['销售误导', '乱收费', '12378'])
        if '物业' in scene_name:
            keywords.extend(['物业不作为', '物业乱收费', '电梯坏了'])
        return list(set(keywords))[:8]

    def _infer_city_code(self, channel):
        """推断城市编码"""
        text = channel.get('name', '') + channel.get('scope', '') + channel.get('regulator', '')
        if '成都' in text:
            return 'chengdu'
        if '四川' in text:
            return 'sichuan'
        return 'national'

    def _infer_channel_type(self, channel):
        """推断渠道类型"""
        name = channel.get('name', '')
        if '平台' in name or '网站' in name or '中心' in name:
            return 'platform'
        if '热线' in name or '电话' in name:
            return 'hotline'
        return 'official'

    def _deduplicate(self):
        """去重（按名称+电话去重）"""
        seen = {}
        unique = []
        for ch in self.channels:
            key = (ch.get('name', ''), ch.get('phone', ''))
            if key in seen:
                # 合并标签
                existing = seen[key]
                existing['tags'] = list(set(existing.get('tags', []) + ch.get('tags', [])))
                self.warnings.append(f"去重: {ch['name']} (重复渠道已合并标签)")
            else:
                seen[key] = ch
                unique.append(ch)
        self.channels = unique
        # 重新分配ID
        for i, ch in enumerate(self.channels, 1):
            ch['id'] = f'ch_{i:03d}'


# ============================================================
# 数据清洗器
# ============================================================

class DataCleaner:
    """数据清洗：去除空格、统一格式、数组拆分"""

    @staticmethod
    def clean_text(text):
        """清洗文本"""
        if not text:
            return ''
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # 去除首尾空格
        text = text.strip()
        # 合并多余空行
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    @staticmethod
    def split_tags(tags_str):
        """将逗号分隔的标签字符串拆分为数组"""
        if not tags_str:
            return []
        if isinstance(tags_str, list):
            return tags_str
        # 支持中英文逗号分隔
        parts = re.split(r'[,，、;；]', tags_str)
        return [p.strip() for p in parts if p.strip()]

    @staticmethod
    def clean_channel(channel):
        """清洗单个渠道"""
        for key in ['name', 'phone', 'phone_note', 'website', 'regulator', 'scope',
                     'precondition', 'legal_basis', 'tips', 'source', 'category_l1', 'category_l2']:
            if key in channel:
                channel[key] = DataCleaner.clean_text(channel[key])

        # tags处理
        if 'tags' in channel and isinstance(channel['tags'], str):
            channel['tags'] = DataCleaner.split_tags(channel['tags'])

        # hot_level确保是int
        if 'hot_level' in channel:
            try:
                channel['hot_level'] = int(channel['hot_level'])
            except (ValueError, TypeError):
                channel['hot_level'] = 3

        return channel

    @staticmethod
    def clean_script(script):
        """清洗单个话术"""
        for key in ['scene_name', 'applicable', 'phone_script', 'written_complainant',
                     'written_respondent', 'written_request', 'written_facts',
                     'written_evidence', 'legal_basis']:
            if key in script:
                script[key] = DataCleaner.clean_text(script[key])

        if 'keywords' in script and isinstance(script['keywords'], str):
            script['keywords'] = DataCleaner.split_tags(script['keywords'])

        if 'evidence_list' in script and isinstance(script['evidence_list'], str):
            script['evidence_list'] = DataCleaner.split_tags(script['evidence_list'])

        return script


# ============================================================
# Excel生成器
# ============================================================

class ExcelWriter:
    """生成Excel维护文件"""

    def __init__(self):
        self.wb = Workbook()
        self._setup_styles()

    def _setup_styles(self):
        """设置样式"""
        self.header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
        self.header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        self.required_fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
        self.warning_fill = PatternFill(start_color='FCE4D6', end_color='FCE4D6', fill_type='solid')
        self.normal_font = Font(name='微软雅黑', size=10)
        self.alignment = Alignment(vertical='top', wrap_text=True)
        self.border = Border(
            left=Side(style='thin', color='D9D9D9'),
            right=Side(style='thin', color='D9D9D9'),
            top=Side(style='thin', color='D9D9D9'),
            bottom=Side(style='thin', color='D9D9D9')
        )

    def write_channels(self, channels):
        """写入渠道表"""
        ws = self.wb.active
        ws.title = '渠道表'

        # 写表头
        for col, (field, label, required, width) in enumerate(CHANNEL_COLUMNS, 1):
            cell = ws.cell(row=1, column=col, value=label)
            cell.font = self.header_font
            cell.fill = self.header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = self.border
            ws.column_dimensions[cell.column_letter].width = width

        # 写数据
        for row_idx, channel in enumerate(channels, 2):
            for col, (field, label, required, width) in enumerate(CHANNEL_COLUMNS, 1):
                value = channel.get(field, '')
                if field == 'tags' and isinstance(value, list):
                    value = ','.join(value)
                if field == 'ext' and isinstance(value, dict):
                    value = json.dumps(value, ensure_ascii=False)
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.font = self.normal_font
                cell.alignment = self.alignment
                cell.border = self.border
                if required and not value:
                    cell.fill = self.warning_fill

        # 冻结首行
        ws.freeze_panes = 'A2'

        # 添加数据验证（枚举字段）
        self._add_data_validation(ws, 'status', VALID_STATUS)
        self._add_data_validation(ws, 'city_code', VALID_CITY_CODE)
        self._add_data_validation(ws, 'channel_type', VALID_CHANNEL_TYPE)

    def write_scripts(self, scripts):
        """写入话术表"""
        ws = self.wb.create_sheet('话术表')

        for col, (field, label, required, width) in enumerate(SCRIPT_COLUMNS, 1):
            cell = ws.cell(row=1, column=col, value=label)
            cell.font = self.header_font
            cell.fill = self.header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = self.border
            ws.column_dimensions[cell.column_letter].width = width

        for row_idx, script in enumerate(scripts, 2):
            for col, (field, label, required, width) in enumerate(SCRIPT_COLUMNS, 1):
                value = script.get(field, '')
                if field in ['keywords', 'evidence_list'] and isinstance(value, list):
                    value = ','.join(value)
                if field == 'ext' and isinstance(value, dict):
                    value = json.dumps(value, ensure_ascii=False)
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.font = self.normal_font
                cell.alignment = self.alignment
                cell.border = self.border
                if required and not value:
                    cell.fill = self.warning_fill

        ws.freeze_panes = 'A2'

    def write_config(self, config):
        """写入配置表（key-value格式）"""
        ws = self.wb.create_sheet('配置表')
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 60
        ws.column_dimensions['C'].width = 30

        headers = ['配置项', '值', '说明']
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = self.header_font
            cell.fill = self.header_fill

        row = 2
        for key, value in config.items():
            ws.cell(row=row, column=1, value=key)
            ws.cell(row=row, column=2, value=json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value))
            row += 1

    def write_data_dictionary(self):
        """写入数据字典"""
        ws = self.wb.create_sheet('数据字典')
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 10
        ws.column_dimensions['E'].width = 40

        headers = ['所属表', '字段名', '中文说明', '必填', '枚举值/格式']
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = self.header_font
            cell.fill = self.header_fill

        row = 2
        for field, label, required, width in CHANNEL_COLUMNS:
            ws.cell(row=row, column=1, value='渠道表')
            ws.cell(row=row, column=2, value=field)
            ws.cell(row=row, column=3, value=label)
            ws.cell(row=row, column=4, value='是' if required else '否')
            if field == 'status':
                ws.cell(row=row, column=5, value=','.join(VALID_STATUS))
            elif field == 'city_code':
                ws.cell(row=row, column=5, value=','.join(VALID_CITY_CODE))
            elif field == 'channel_type':
                ws.cell(row=row, column=5, value=','.join(VALID_CHANNEL_TYPE))
            elif field == 'hot_level':
                ws.cell(row=row, column=5, value='1-5整数')
            row += 1

    def write_update_log(self):
        """写入更新日志"""
        ws = self.wb.create_sheet('更新日志')
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 40
        ws.column_dimensions['D'].width = 15

        headers = ['版本', '日期', '变更内容', '操作人']
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = self.header_font
            cell.fill = self.header_fill

        ws.cell(row=2, column=1, value='2026.08')
        ws.cell(row=2, column=2, value=datetime.now().strftime('%Y-%m-%d'))
        ws.cell(row=2, column=3, value='首次提取，包含109个行业渠道、16个高层级平台、5套话术模板')

    def _add_data_validation(self, ws, field_name, valid_values):
        """添加数据验证（下拉选择）"""
        # 找到字段所在列
        col_idx = None
        for col, (f, label, req, width) in enumerate(CHANNEL_COLUMNS, 1):
            if f == field_name:
                col_idx = col
                break
        if not col_idx:
            return

        col_letter = ws.cell(row=1, column=col_idx).column_letter
        dv = DataValidation(type='list', formula1=f'"{",".join(valid_values)}"', allow_blank=True)
        dv.error = f'请选择: {",".join(valid_values)}'
        dv.errorTitle = '输入错误'
        ws.add_data_validation(dv)
        dv.add(f'{col_letter}2:{col_letter}1000')

    def save(self, filepath):
        """保存Excel文件"""
        os.makedirs(os.path.dirname(filepath) or '.', exist_ok=True)
        self.wb.save(filepath)
        logger.info(f"Excel已保存: {filepath}")


# ============================================================
# Excel读取器
# ============================================================

class ExcelReader:
    """读取Excel维护文件"""

    @staticmethod
    def read_channels(filepath):
        """读取渠道表"""
        wb = load_workbook(filepath, data_only=True)
        ws = wb['渠道表']

        # 读取表头，建立字段名映射
        headers = {}
        for col in range(1, ws.max_column + 1):
            label = ws.cell(row=1, column=col).value
            if label:
                # 从CHANNEL_COLUMNS查找字段名
                for field, l, req, width in CHANNEL_COLUMNS:
                    if l == label:
                        headers[col] = field
                        break

        channels = []
        for row in range(2, ws.max_row + 1):
            channel = {}
            has_data = False
            for col, field in headers.items():
                value = ws.cell(row=row, column=col).value
                if value is not None and str(value).strip():
                    has_data = True
                channel[field] = value if value is not None else ''

            if has_data and channel.get('name'):
                # tags字符串转数组
                if 'tags' in channel and isinstance(channel['tags'], str):
                    channel['tags'] = DataCleaner.split_tags(channel['tags'])
                channels.append(DataCleaner.clean_channel(channel))

        return channels

    @staticmethod
    def read_scripts(filepath):
        """读取话术表"""
        wb = load_workbook(filepath, data_only=True)
        if '话术表' not in wb.sheetnames:
            return []

        ws = wb['话术表']
        headers = {}
        for col in range(1, ws.max_column + 1):
            label = ws.cell(row=1, column=col).value
            if label:
                for field, l, req, width in SCRIPT_COLUMNS:
                    if l == label:
                        headers[col] = field
                        break

        scripts = []
        for row in range(2, ws.max_row + 1):
            script = {}
            has_data = False
            for col, field in headers.items():
                value = ws.cell(row=row, column=col).value
                if value is not None and str(value).strip():
                    has_data = True
                script[field] = value if value is not None else ''

            if has_data and script.get('scene_name'):
                if 'keywords' in script and isinstance(script['keywords'], str):
                    script['keywords'] = DataCleaner.split_tags(script['keywords'])
                if 'evidence_list' in script and isinstance(script['evidence_list'], str):
                    script['evidence_list'] = DataCleaner.split_tags(script['evidence_list'])
                scripts.append(DataCleaner.clean_script(script))

        return scripts


# ============================================================
# 法律法规提取器（去重存储，场景引用ID）
# ============================================================

class LawExtractor:
    """
    从渠道/话术的legal_basis字段中提取法律法规，去重存储
    解决上千场景反复引用同一部法律导致的体积膨胀问题
    场景表只存 law_ids: ["law_001", "law_023"]，不存法律全文
    """

    # 法律法规名称匹配模式
    LAW_PATTERNS = [
        r'《([^》]+(?:法|条例|办法|规定|细则|司法解释|指导意见|通知|公告))》',
        r'(中华人民共和国[^，。；、\s]{2,20}(?:法|条例|规定))',
        r'(?:根据|依据|按照)([^，。；、\s]{2,30}(?:法|条例|办法|规定|细则))',
    ]

    # 常见法律简称映射
    LAW_ALIASES = {
        '消法': '中华人民共和国消费者权益保护法',
        '消费者权益保护法': '中华人民共和国消费者权益保护法',
        '电商法': '中华人民共和国电子商务法',
        '电子商务法': '中华人民共和国电子商务法',
        '劳动合同法': '中华人民共和国劳动合同法',
        '劳动法': '中华人民共和国劳动法',
        '社会保险法': '中华人民共和国社会保险法',
        '邮政法': '中华人民共和国邮政法',
        '电信条例': '中华人民共和国电信条例',
        '物业管理条例': '物业管理条例',
        '价格法': '中华人民共和国价格法',
        '反垄断法': '中华人民共和国反垄断法',
        '反不正当竞争法': '中华人民共和国反不正当竞争法',
        '食品安全法': '中华人民共和国食品安全法',
        '药品管理法': '中华人民共和国药品管理法',
        '环境保护法': '中华人民共和国环境保护法',
        '噪声污染防治法': '中华人民共和国噪声污染防治法',
        '个人信息保护法': '中华人民共和国个人信息保护法',
        '数据安全法': '中华人民共和国数据安全法',
        '网络安全法': '中华人民共和国网络安全法',
        '民法典': '中华人民共和国民法典',
        '民事诉讼法': '中华人民共和国民事诉讼法',
        '行政复议法': '中华人民共和国行政复议法',
        '行政诉讼法': '中华人民共和国行政诉讼法',
        '银行业监督管理法': '中华人民共和国银行业监督管理法',
        '保险法': '中华人民共和国保险法',
        '证券法': '中华人民共和国证券法',
        '教育法': '中华人民共和国教育法',
        '义务教育法': '中华人民共和国义务教育法',
        '道路交通安全法': '中华人民共和国道路交通安全法',
        '消防法': '中华人民共和国消防法',
        '安全生产法': '中华人民共和国安全生产法',
        '税收征收管理法': '中华人民共和国税收征收管理法',
        '土地管理法': '中华人民共和国土地管理法',
        '城乡规划法': '中华人民共和国城乡规划法',
        '建筑法': '中华人民共和国建筑法',
        '广告法': '中华人民共和国广告法',
        '产品质量法': '中华人民共和国产品质量法',
        '旅游法': '中华人民共和国旅游法',
        '传染病防治法': '中华人民共和国传染病防治法',
        '基本医疗卫生与健康促进法': '中华人民共和国基本医疗卫生与健康促进法',
    }

    def __init__(self):
        self.laws = {}  # law_id -> {id, name, full_name, article, source_count}
        self.name_to_id = {}  # 法律名称 -> law_id
        self._counter = 0

    def extract_from_documents(self, channels, scripts):
        """从所有文档中提取法律法规，并更新文档的law_ids字段"""
        logger.info("开始提取法律法规...")

        all_docs = []
        for ch in channels:
            all_docs.append(('channel', ch))
        for s in scripts:
            all_docs.append(('script', s))

        for doc_type, doc in all_docs:
            legal_text = doc.get('legal_basis', '') or ''
            if not legal_text:
                continue

            law_ids = []
            # 提取法律名称
            extracted_names = self._extract_law_names(legal_text)

            for name in extracted_names:
                # 标准化名称
                standard_name = self._standardize_name(name)
                if not standard_name:
                    continue

                # 去重：同一文档中同一法律只引用一次
                if standard_name in self.name_to_id:
                    law_id = self.name_to_id[standard_name]
                    self.laws[law_id]['source_count'] += 1
                else:
                    self._counter += 1
                    law_id = f'law_{self._counter:03d}'
                    self.name_to_id[standard_name] = law_id
                    self.laws[law_id] = {
                        'id': law_id,
                        'name': standard_name,
                        'full_name': standard_name,
                        'article': self._extract_article(legal_text, standard_name),
                        'source_count': 1,
                    }

                if law_id not in law_ids:
                    law_ids.append(law_id)

            # 更新文档的law_ids字段
            if law_ids:
                doc['law_ids'] = ','.join(law_ids)
            else:
                doc['law_ids'] = ''

        logger.info(f"法律法规提取完成: {len(self.laws)}部不同法律")
        return list(self.laws.values())

    def _extract_law_names(self, text):
        """从文本中提取法律法规名称"""
        names = []
        for pattern in self.LAW_PATTERNS:
            matches = re.findall(pattern, text)
            names.extend(matches)

        # 去重，保持顺序
        seen = set()
        result = []
        for name in names:
            name = name.strip()
            if name and name not in seen:
                seen.add(name)
                result.append(name)
        return result

    def _standardize_name(self, name):
        """标准化法律名称（处理简称、空格、标点）"""
        name = name.strip().strip('《》').strip()

        # 处理简称
        if name in self.LAW_ALIASES:
            return self.LAW_ALIASES[name]

        # 检查是否包含已知法律的关键词
        for alias, full in self.LAW_ALIASES.items():
            if alias in name and len(name) < len(full) + 5:
                return full

        # 太短的可能不是法律名称
        if len(name) < 4:
            return None

        return name

    def _extract_article(self, text, law_name):
        """提取法律条文（如"第X条"）"""
        # 在法律名称附近查找"第X条"
        pattern = re.escape(law_name) + r'[^第]{0,20}(第[一二三四五六七八九十百千\d]+条)'
        match = re.search(pattern, text)
        if match:
            return match.group(1)
        # 全局查找第一个"第X条"
        match = re.search(r'(第[一二三四五六七八九十百千\d]+条)', text)
        if match:
            return match.group(1)
        return ''

    def get_laws_list(self):
        """获取法律法规列表（按引用次数排序）"""
        return sorted(self.laws.values(), key=lambda x: x['source_count'], reverse=True)


# ============================================================
# 预构建索引生成器（搜索加速核心）
# ============================================================

class IndexBuilder:
    """
    预构建倒排索引和搜索联想Trie树
    在构建阶段生成索引文件，小程序运行时直接加载，无需现场构建
    相比运行时构建：省掉200ms启动时间，分词质量更高（Python端可做更复杂处理）
    """

    # 字段权重配置（与小程序端保持一致）
    FIELD_WEIGHTS = {
        'name': 100,
        'phone': 90,
        'tags': 80,
        'scope': 50,
        'regulator': 40,
        'legal_basis': 30,
        'tips': 25,
        'scene_name': 80,
        'keywords': 70,
        'applicable': 40,
        'phone_script': 30,
    }

    def __init__(self):
        self.dictionary = set()  # 全部term集合
        self.postings = defaultdict(list)  # term -> [{doc, type, field, weight, pos}]
        self.doc_length = {}  # docId -> 文档长度（term数）
        self.trie = {}  # 搜索联想Trie树
        self._custom_dict = set()  # 自定义分词词典

    def build(self, channels, scripts):
        """构建完整索引"""
        logger.info("开始构建预搜索索引...")

        # 1. 构建自定义分词词典（从名称、标签、分类中提取）
        self._build_dictionary(channels, scripts)

        # 2. 索引渠道数据
        for ch in channels:
            self._index_document(ch['id'], 'channel', ch)

        # 3. 索引话术数据
        for s in scripts:
            self._index_document(s['id'], 'script', s)

        # 4. 构建搜索联想Trie树
        self._build_trie(channels, scripts)

        # 5. 计算统计信息
        total_docs = len(channels) + len(scripts)
        avg_doc_length = sum(self.doc_length.values()) / max(total_docs, 1)

        # 6. 压缩索引格式（用数字索引代替字符串，大幅减小体积）
        # doc_id -> 数字索引
        all_doc_ids = list(self.doc_length.keys())
        doc_map = {doc_id: idx for idx, doc_id in enumerate(all_doc_ids)}
        # field_name -> 数字索引
        field_names = list(self.FIELD_WEIGHTS.keys())
        field_map = {name: idx for idx, name in enumerate(field_names)}
        # type -> 数字
        type_map = {'channel': 0, 'script': 1}

        # 压缩postings: term -> [[doc_idx, field_idx, type_idx, weight], ...]
        compressed_postings = {}
        for term, posting_list in self.postings.items():
            # 同一term同一文档同一field只保留一条（取最大weight）
            merged = {}
            for p in posting_list:
                key = (doc_map[p['doc']], field_map[p['field']], type_map[p['type']])
                if key not in merged or p['w'] > merged[key]:
                    merged[key] = p['w']
            compressed_postings[term] = [[k[0], k[1], k[2], v] for k, v in merged.items()]

        # 压缩doc_length: 用数字索引作为key
        compressed_doc_length = {str(doc_map[doc_id]): length for doc_id, length in self.doc_length.items()}

        index_data = {
            'v': DEFAULT_CONFIG['data_version'],
            'built_at': datetime.now().isoformat(),
            'total_docs': total_docs,
            'term_count': len(self.dictionary),
            'avg_dl': round(avg_doc_length, 1),
            'doc_map': all_doc_ids,  # 数字索引 -> doc_id
            'field_map': field_names,  # 数字索引 -> field_name
            'field_weights': [self.FIELD_WEIGHTS[f] for f in field_names],
            'doc_length': compressed_doc_length,
            'postings': compressed_postings,
        }

        trie_data = {
            'v': DEFAULT_CONFIG['data_version'],
            'trie': self.trie,
            'hot': DEFAULT_CONFIG['hot_search_words'],
        }

        logger.info(f"索引构建完成: {len(self.dictionary)}个term, {total_docs}个文档")
        return index_data, trie_data

    def _build_dictionary(self, channels, scripts):
        """从数据中提取分词词典"""
        for ch in channels:
            # 名称中的2-6字词
            for word in re.findall(r'[\u4e00-\u9fa5]{2,6}', ch.get('name', '')):
                self._custom_dict.add(word)
            # 标签
            for tag in ch.get('tags', []):
                if isinstance(tag, str) and 2 <= len(tag) <= 8:
                    self._custom_dict.add(tag)
            # 分类名
            for cat in [ch.get('category_l1', ''), ch.get('category_l2', '')]:
                for word in re.findall(r'[\u4e00-\u9fa5]{2,6}', cat):
                    self._custom_dict.add(word)

        for s in scripts:
            for kw in s.get('keywords', []):
                if isinstance(kw, str) and 2 <= len(kw) <= 8:
                    self._custom_dict.add(kw)
            for word in re.findall(r'[\u4e00-\u9fa5]{2,6}', s.get('scene_name', '')):
                self._custom_dict.add(word)

        # 加入高频维权词
        self._custom_dict.update([
            '投诉', '举报', '申诉', '维权', '纠纷', '违法', '违规', '收费',
            '退款', '丢失', '破损', '延误', '诈骗', '欺诈', '误导', '不作为',
            '乱收费', '欠薪', '社保', '公积金', '物业', '快递', '电信', '银行',
            '保险', '医疗', '教育', '食品', '药品', '环保', '劳动', '税务',
            '公安', '市场监管', '消费者', '个人信息', '噪音', '扰民', '假货',
        ])

        logger.debug(f"分词词典: {len(self._custom_dict)}个词")

    def _index_document(self, doc_id, doc_type, doc):
        """索引单个文档"""
        doc_term_count = 0

        for field, weight in self.FIELD_WEIGHTS.items():
            text = doc.get(field, '')
            if not text:
                continue
            if isinstance(text, list):
                text = ' '.join(str(t) for t in text)

            terms = self._tokenize(str(text))
            for pos, term in enumerate(terms):
                self.dictionary.add(term)
                self.postings[term].append({
                    'doc': doc_id,
                    'type': doc_type,
                    'field': field,
                    'w': weight,
                    'p': pos,
                })
                doc_term_count += 1

        self.doc_length[doc_id] = doc_term_count

    def _tokenize(self, text):
        """
        中文分词（基于词典的正向最大匹配 + 单字兜底 + 英文数字）
        比小程序端的简单分词质量更高
        """
        terms = []
        i = 0
        while i < len(text):
            ch = text[i]

            # 英文/数字连续字符
            if re.match(r'[a-zA-Z0-9]', ch):
                j = i
                while j < len(text) and re.match(r'[a-zA-Z0-9]', text[j]):
                    j += 1
                word = text[i:j].lower()
                if len(word) >= 2:
                    terms.append(word)
                # 数字也单独加入（电话号码搜索）
                if word.isdigit() and len(word) >= 3:
                    terms.append(word)
                i = j
                continue

            # 中文：正向最大匹配（最长6字）
            if '\u4e00' <= ch <= '\u9fff':
                matched = False
                for length in range(min(6, len(text) - i), 1, -1):
                    word = text[i:i + length]
                    if word in self._custom_dict:
                        terms.append(word)
                        i += length
                        matched = True
                        break
                if not matched:
                    # 单字兜底（保证召回率）
                    terms.append(ch)
                    i += 1
                continue

            # 其他字符跳过
            i += 1

        return terms

    def _build_trie(self, channels, scripts):
        """构建搜索联想Trie树"""
        # 收集所有可联想的词
        all_words = set()

        for ch in channels:
            all_words.add(ch['name'])
            for tag in ch.get('tags', []):
                if isinstance(tag, str) and 2 <= len(tag) <= 10:
                    all_words.add(tag)

        for s in scripts:
            all_words.add(s['scene_name'])
            for kw in s.get('keywords', []):
                if isinstance(kw, str) and 2 <= len(kw) <= 10:
                    all_words.add(kw)

        # 加入热门搜索词
        all_words.update(DEFAULT_CONFIG['hot_search_words'])

        # 构建Trie（每个节点存以该前缀开头的词列表）
        for word in all_words:
            if not word or len(word) < 2:
                continue
            node = self.trie
            for ch in word:
                if ch not in node:
                    node[ch] = {'$': []}
                node = node[ch]
                # 每个节点都存完整词列表（用于前缀联想）
                if '$' not in node:
                    node['$'] = []
                if word not in node['$'] and len(node['$']) < 10:
                    node['$'].append(word)


# ============================================================
# JSON生成器
# ============================================================

class JsonWriter:
    """生成JSON运行时文件"""

    @staticmethod
    def write_all(channels, scripts, platforms, output_dir, minify=False, shard_threshold=300):
        """
        生成所有JSON文件
        :param shard_threshold: 渠道数超过此值时自动按一级分类分片输出
        """
        os.makedirs(output_dir, exist_ok=True)

        files = {}

        # 1. 提取法律法规（去重存储，场景引用ID）
        law_extractor = LawExtractor()
        laws = law_extractor.extract_from_documents(channels, scripts)
        files['laws.json'] = laws

        # 2. 渠道数据（超过阈值时分片）
        if len(channels) > shard_threshold:
            logger.info(f"渠道数{len(channels)}超过分片阈值{shard_threshold}，按一级分类分片输出")
            shard_map = defaultdict(list)
            for ch in channels:
                cat = ch.get('category_l1', '未分类')
                # 生成安全的文件名
                safe_cat = re.sub(r'[^\w\u4e00-\u9fa5]', '_', cat)[:20]
                shard_map[safe_cat].append(ch)

            shard_index = []
            for idx, (cat_name, cat_channels) in enumerate(sorted(shard_map.items()), 1):
                filename = f'channels_{idx:02d}_{cat_name}.json'
                files[filename] = cat_channels
                shard_index.append({
                    'shard': filename,
                    'category': cat_name,
                    'count': len(cat_channels),
                })
            files['channels_index.json'] = shard_index
            logger.info(f"渠道分片: {len(shard_index)}个分片")
        else:
            files['channels.json'] = channels

        # scripts.json
        files['scripts.json'] = scripts

        # platforms.json
        files['platforms.json'] = platforms

        # config.json
        files['config.json'] = DEFAULT_CONFIG

        # categories.json（从渠道数据提取）
        categories = JsonWriter._extract_categories(channels)
        files['categories.json'] = categories

        # 预构建搜索索引（倒排索引 + Trie联想树）
        # 注意：索引只包含term→docId映射，不包含详情数据，体积可控
        index_builder = IndexBuilder()
        search_index, suggest_trie = index_builder.build(channels, scripts)
        files['search_index.json'] = search_index
        files['suggest_trie.json'] = suggest_trie

        # 写入文件
        manifest = []
        for filename, data in files.items():
            filepath = os.path.join(output_dir, filename)
            indent = None if minify else 2
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=indent)

            # 计算校验和
            with open(filepath, 'rb') as f:
                checksum = hashlib.sha256(f.read()).hexdigest()[:16]
            size = os.path.getsize(filepath)
            manifest.append({'file': filename, 'size': size, 'sha256': checksum})
            logger.info(f"JSON已生成: {filename} ({size} bytes)")

        # 生成微信小程序专用分片JS数据（索引+按需加载分片）
        # 解决单文件500KB限制，支持扩展到几千条数据
        shard_info = JsonWriter.write_miniprogram_shards(channels, output_dir)
        logger.info(f"小程序分片生成完成: {shard_info['num_parts']}个分片")

        # _manifest.json
        manifest_data = {
            'version': DEFAULT_CONFIG['data_version'],
            'generated_at': datetime.now().isoformat(),
            'channel_count': len(channels),
            'script_count': len(scripts),
            'platform_count': len(platforms),
            'law_count': len(laws),
            'is_sharded': len(channels) > shard_threshold,
            'miniprogram_shards': shard_info,
            'files': manifest
        }
        manifest_path = os.path.join(output_dir, '_manifest.json')
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump(manifest_data, f, ensure_ascii=False, indent=2)
        logger.info(f"清单已生成: _manifest.json")

        return manifest_data

    @staticmethod
    def write_miniprogram_shards(channels, output_dir, channels_per_part=50):
        """
        生成微信小程序专用的分片JS数据文件（索引+按需加载分片）
        解决单文件500KB限制，支持扩展到几千条数据

        生成文件：
        - channels_index.js: 轻量索引（id/name/phone/tags/part_num），启动时加载
        - channels_part_N.js: 详细数据分片，点击详情时按需加载
        - channels_config.js: 分片配置信息

        :param channels: 渠道数据列表
        :param output_dir: 输出目录
        :param channels_per_part: 每个分片的渠道数量（默认50条，约60KB）
        :return: 分片统计信息
        """
        os.makedirs(output_dir, exist_ok=True)
        total = len(channels)
        num_parts = (total + channels_per_part - 1) // channels_per_part

        logger.info(f"生成小程序分片数据: {total}条渠道, {num_parts}个分片, 每片{channels_per_part}条")

        # 1. 生成轻量索引
        index_data = []
        for i, ch in enumerate(channels):
            part_num = (i // channels_per_part) + 1
            index_data.append({
                'id': ch.get('id', ''),
                'name': ch.get('name', ''),
                'category_l1': ch.get('category_l1', ''),
                'category_l2': ch.get('category_l2', ''),
                'tags': ch.get('tags', []),
                'hot_level': ch.get('hot_level', 0),
                'phone': ch.get('phone', ''),
                'part_num': part_num
            })

        index_path = os.path.join(output_dir, 'channels_index.js')
        with open(index_path, 'w', encoding='utf-8') as f:
            f.write(f"// channels_index.js - 渠道轻量索引（启动时加载，用于列表展示和搜索）\n")
            f.write(f"// 由convert.py自动生成，共{len(index_data)}条记录\n")
            f.write(f"module.exports = {json.dumps(index_data, ensure_ascii=False, indent=2)};\n")
        logger.info(f"  索引文件: channels_index.js ({os.path.getsize(index_path)//1024}KB)")

        # 2. 生成分片详细数据
        part_files = []
        for part_num in range(1, num_parts + 1):
            start = (part_num - 1) * channels_per_part
            end = min(start + channels_per_part, total)
            part_channels = channels[start:end]

            part_filename = f'channels_part_{part_num}.js'
            part_path = os.path.join(output_dir, part_filename)
            with open(part_path, 'w', encoding='utf-8') as f:
                f.write(f"// channels_part_{part_num}.js - 渠道数据分片{part_num}/{num_parts}（第{start+1}-{end}条）\n")
                f.write(f"// 由convert.py自动生成，点击详情时按需加载，请勿手动编辑\n")
                f.write(f"module.exports = {json.dumps(part_channels, ensure_ascii=False, indent=2)};\n")

            part_files.append(part_filename)
            logger.info(f"  分片{part_num}: {part_filename} ({len(part_channels)}条, {os.path.getsize(part_path)//1024}KB)")

        # 3. 生成分片配置
        config_data = {
            'total': total,
            'num_parts': num_parts,
            'channels_per_part': channels_per_part,
            'part_files': part_files,
            'index_file': 'channels_index.js',
            'generated_at': datetime.now().isoformat()
        }
        config_path = os.path.join(output_dir, 'channels_config.js')
        with open(config_path, 'w', encoding='utf-8') as f:
            f.write(f"// channels_config.js - 渠道分片配置\n")
            f.write(f"module.exports = {json.dumps(config_data, ensure_ascii=False, indent=2)};\n")
        logger.info(f"  配置文件: channels_config.js")

        return {
            'total': total,
            'num_parts': num_parts,
            'channels_per_part': channels_per_part,
            'index_size': os.path.getsize(index_path),
            'part_files': part_files
        }

    @staticmethod
    def _extract_categories(channels):
        """从渠道数据提取分类树"""
        cat_map = defaultdict(lambda: {'name': '', 'children': set()})
        for ch in channels:
            l1 = ch.get('category_l1', '未分类')
            l2 = ch.get('category_l2', '未分类')
            cat_map[l1]['name'] = l1
            cat_map[l1]['children'].add(l2)

        result = []
        for l1, data in sorted(cat_map.items()):
            result.append({
                'id': f'cat_{len(result)+1:02d}',
                'name': l1,
                'children': [{'name': c} for c in sorted(data['children'])]
            })
        return result


# ============================================================
# 数据校验器
# ============================================================

class Validator:
    """数据校验"""

    def __init__(self):
        self.errors = []
        self.warnings = []

    def validate_channels(self, channels):
        """校验渠道数据"""
        logger.info(f"校验渠道数据: {len(channels)}条")

        # ID唯一性
        ids = [ch['id'] for ch in channels if ch.get('id')]
        dup_ids = [id for id in set(ids) if ids.count(id) > 1]
        if dup_ids:
            self.errors.append(f"渠道ID重复: {dup_ids}")

        for ch in channels:
            cid = ch.get('id', '未知')

            # 必填字段
            for field in ['id', 'name', 'scope', 'category_l1', 'category_l2', 'status', 'hot_level']:
                if not ch.get(field):
                    self.errors.append(f"渠道{cid}缺少必填字段: {field}")

            # 枚举值
            if ch.get('status') and ch['status'] not in VALID_STATUS:
                self.errors.append(f"渠道{cid}status值非法: {ch['status']}")
            if ch.get('city_code') and ch['city_code'] not in VALID_CITY_CODE:
                self.warnings.append(f"渠道{cid}city_code值不在预设枚举: {ch['city_code']}")
            if ch.get('channel_type') and ch['channel_type'] not in VALID_CHANNEL_TYPE:
                self.warnings.append(f"渠道{cid}channel_type值不在预设枚举: {ch['channel_type']}")

            # 标签数量
            tags = ch.get('tags', [])
            if isinstance(tags, list) and len(tags) < 3:
                self.warnings.append(f"渠道{cid}标签不足3个: {tags}")

            # 热度范围
            hot = ch.get('hot_level')
            if hot is not None:
                try:
                    hot_int = int(hot)
                    if hot_int < 1 or hot_int > 5:
                        self.warnings.append(f"渠道{cid}hot_level超出1-5范围: {hot}")
                except (ValueError, TypeError):
                    self.warnings.append(f"渠道{cid}hot_level不是整数: {hot}")

            # 电话格式（简单校验）
            phone = ch.get('phone', '')
            if phone and not re.match(r'^[\d\-+\(\)（）\s,，]+$', str(phone)):
                self.warnings.append(f"渠道{cid}电话格式可能异常: {phone}")

            # 网址格式
            website = ch.get('website', '')
            if website and not website.startswith(('http://', 'https://', 'www.')):
                self.warnings.append(f"渠道{cid}网址格式可能异常: {website}")

        # 外键校验（merged_to, related_script_id）
        channel_ids = set(ids)
        for ch in channels:
            merged = ch.get('merged_to')
            if merged and merged not in channel_ids:
                self.errors.append(f"渠道{ch.get('id')}的merged_to指向不存在的渠道: {merged}")

        return len(self.errors) == 0

    def validate_scripts(self, scripts):
        """校验话术数据"""
        logger.info(f"校验话术数据: {len(scripts)}条")

        ids = [s['id'] for s in scripts if s.get('id')]
        dup_ids = [id for id in set(ids) if ids.count(id) > 1]
        if dup_ids:
            self.errors.append(f"话术ID重复: {dup_ids}")

        for s in scripts:
            sid = s.get('id', '未知')
            for field in ['id', 'scene_name', 'script_type']:
                if not s.get(field):
                    self.errors.append(f"话术{sid}缺少必填字段: {field}")

            # 电话版和书面版至少有一个
            if not s.get('phone_script') and not s.get('written_facts'):
                self.errors.append(f"话术{sid}电话版和书面版不能同时为空")

        return len(self.errors) == 0

    def get_report(self):
        """生成校验报告"""
        report = []
        report.append(f"=== 数据校验报告 ===")
        report.append(f"校验时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"错误数: {len(self.errors)}")
        report.append(f"警告数: {len(self.warnings)}")
        report.append("")

        if self.errors:
            report.append("【错误】")
            for e in self.errors:
                report.append(f"  ✗ {e}")
            report.append("")

        if self.warnings:
            report.append("【警告】")
            for w in self.warnings[:20]:  # 最多显示20条
                report.append(f"  ⚠ {w}")
            if len(self.warnings) > 20:
                report.append(f"  ... 还有{len(self.warnings)-20}条警告")

        report.append("")
        if self.errors:
            report.append("结论: ❌ 校验未通过，请修正错误后重新构建")
        else:
            report.append("结论: ✅ 校验通过" + (f"（{len(self.warnings)}条警告建议关注）" if self.warnings else ""))

        return '\n'.join(report)


# ============================================================
# Diff比较器（增量更新）
# ============================================================

class DiffComparator:
    """比较新旧数据，输出diff报告"""

    @staticmethod
    def compare(old_channels, new_channels):
        """比较渠道数据"""
        old_map = {ch['id']: ch for ch in old_channels}
        new_map = {ch['id']: ch for ch in new_channels}

        added = []
        deleted = []
        modified = []

        # 新增
        for id, ch in new_map.items():
            if id not in old_map:
                added.append(ch)

        # 删除
        for id, ch in old_map.items():
            if id not in new_map:
                deleted.append(ch)

        # 修改
        for id in old_map.keys() & new_map.keys():
            old = old_map[id]
            new = new_map[id]
            changes = {}
            for key in set(list(old.keys()) + list(new.keys())):
                old_val = old.get(key, '')
                new_val = new.get(key, '')
                if old_val != new_val:
                    changes[key] = {'old': old_val, 'new': new_val}
            if changes:
                modified.append({'id': id, 'name': old.get('name', ''), 'changes': changes})

        return {
            'added': added,
            'deleted': deleted,
            'modified': modified,
            'old_count': len(old_channels),
            'new_count': len(new_channels)
        }

    @staticmethod
    def generate_report(diff):
        """生成diff报告"""
        report = []
        report.append("=== 数据增量更新Diff报告 ===")
        report.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"旧数据: {diff['old_count']}条 → 新数据: {diff['new_count']}条")
        report.append(f"新增: {len(diff['added'])}条, 删除: {len(diff['deleted'])}条, 修改: {len(diff['modified'])}条")
        report.append("")

        if diff['added']:
            report.append("【新增渠道】")
            for ch in diff['added']:
                report.append(f"  + {ch['id']} {ch.get('name', '无名称')}")
            report.append("")

        if diff['deleted']:
            report.append("【删除渠道】（建议软删除，不物理移除）")
            for ch in diff['deleted']:
                report.append(f"  - {ch['id']} {ch.get('name', '无名称')}")
            report.append("")

        if diff['modified']:
            report.append("【修改渠道】")
            for m in diff['modified']:
                report.append(f"  ~ {m['id']} {m['name']}")
                for field, change in m['changes'].items():
                    old = str(change['old'])[:50]
                    new = str(change['new'])[:50]
                    report.append(f"    {field}: {old} → {new}")
            report.append("")

        report.append("建议:")
        report.append("  1. 新增渠道: 确认后自动加入Excel")
        report.append("  2. 删除渠道: 建议设置status=discontinued，保留ID引用完整性")
        report.append("  3. 修改渠道: 逐字段确认，保留人工修改过的字段")
        report.append("  4. 冲突字段: 需人工逐条确认采用哪个值")

        return '\n'.join(report)


# ============================================================
# 主函数
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='维权通数据转换工具: Word → Excel → JSON')
    parser.add_argument('--input', required=True, help='输入路径（Word目录或Excel文件）')
    parser.add_argument('--output', required=True, help='输出路径（目录或文件）')
    parser.add_argument('--mode', choices=['extract', 'build', 'all', 'validate'], default='all',
                        help='运行模式: extract(Word→Excel), build(Excel→JSON), all(一键), validate(仅校验)')
    parser.add_argument('--existing', help='现有Excel文件路径（用于diff增量更新）')
    parser.add_argument('--diff', action='store_true', help='启用diff增量更新模式')
    parser.add_argument('--minify', action='store_true', help='生成压缩版JSON（无空格换行）')
    parser.add_argument('--verbose', action='store_true', help='详细日志输出')

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    logger.info("=" * 60)
    logger.info("维权通数据转换工具")
    logger.info(f"模式: {args.mode}, 输入: {args.input}, 输出: {args.output}")
    logger.info("=" * 60)

    # 模式1: 仅校验
    if args.mode == 'validate':
        if not args.input.endswith('.xlsx'):
            logger.error("validate模式需要输入Excel文件")
            exit(1)
        channels = ExcelReader.read_channels(args.input)
        scripts = ExcelReader.read_scripts(args.input)
        validator = Validator()
        validator.validate_channels(channels)
        validator.validate_scripts(scripts)
        report = validator.get_report()
        print(report)
        report_path = os.path.join(os.path.dirname(args.input) or '.', f'校验报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt')
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report)
        logger.info(f"报告已保存: {report_path}")
        return

    # 模式2: Word → Excel
    if args.mode in ['extract', 'all']:
        parser_obj = WordParser()
        if os.path.isdir(args.input):
            parser_obj.parse_directory(args.input)
        elif args.input.endswith('.docx'):
            parser_obj.parse_file(args.input)
        else:
            logger.error(f"不支持的输入格式: {args.input}")
            exit(1)

        # 清洗数据
        channels = [DataCleaner.clean_channel(ch) for ch in parser_obj.channels]
        scripts = [DataCleaner.clean_script(s) for s in parser_obj.scripts]
        platforms = []  # 平台数据暂从渠道中分离（含"平台"类型的渠道）

        # Diff增量更新
        if args.diff and args.existing and os.path.exists(args.existing):
            logger.info("启用Diff增量更新模式")
            old_channels = ExcelReader.read_channels(args.existing)
            diff = DiffComparator.compare(old_channels, channels)
            diff_report = DiffComparator.generate_report(diff)
            print(diff_report)
            diff_path = os.path.join(os.path.dirname(args.output) or '.', f'diff报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt')
            with open(diff_path, 'w', encoding='utf-8') as f:
                f.write(diff_report)
            logger.info(f"Diff报告已保存: {diff_path}")

        # 生成Excel
        excel_path = args.output if args.output.endswith('.xlsx') else os.path.join(args.output, '维权通数据_维护版.xlsx')
        writer = ExcelWriter()
        writer.write_channels(channels)
        writer.write_scripts(scripts)
        writer.write_config(DEFAULT_CONFIG)
        writer.write_data_dictionary()
        writer.write_update_log()
        writer.save(excel_path)

        # 输出提取警告
        if parser_obj.warnings:
            logger.info(f"提取警告({len(parser_obj.warnings)}条):")
            for w in parser_obj.warnings[:10]:
                logger.info(f"  ⚠ {w}")
            if len(parser_obj.warnings) > 10:
                logger.info(f"  ... 还有{len(parser_obj.warnings)-10}条警告")

        # 如果是extract模式，到此结束
        if args.mode == 'extract':
            logger.info("提取完成！")
            return

        # all模式继续构建JSON，使用刚生成的Excel
        args.input = excel_path

    # 模式3: Excel → JSON
    if args.mode in ['build', 'all']:
        if not args.input.endswith('.xlsx'):
            logger.error("build模式需要输入Excel文件")
            exit(1)

        channels = ExcelReader.read_channels(args.input)
        scripts = ExcelReader.read_scripts(args.input)
        platforms = []

        logger.info(f"读取Excel: 渠道{len(channels)}条, 话术{len(scripts)}条")

        # 校验
        validator = Validator()
        channel_ok = validator.validate_channels(channels)
        script_ok = validator.validate_scripts(scripts)
        report = validator.get_report()
        print(report)

        if not channel_ok or not script_ok:
            logger.error("数据校验未通过，终止构建。请修正错误后重新运行。")
            report_path = os.path.join(args.output if os.path.isdir(args.output) else '.', f'校验报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt')
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(report)
            exit(1)

        # 生成JSON
        output_dir = args.output if os.path.isdir(args.output) else os.path.dirname(args.output)
        manifest = JsonWriter.write_all(channels, scripts, platforms, output_dir, minify=args.minify)

        logger.info("=" * 60)
        logger.info("构建完成！")
        logger.info(f"数据版本: {manifest['version']}")
        logger.info(f"渠道: {manifest['channel_count']}条, 话术: {manifest['script_count']}条, 平台: {manifest['platform_count']}条")
        logger.info(f"输出目录: {output_dir}")
        logger.info("=" * 60)


if __name__ == '__main__':
    main()
